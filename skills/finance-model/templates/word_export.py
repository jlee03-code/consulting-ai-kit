"""
BCG research output -> Word (.docx) exporter.
Sources rendered as clickable hyperlinks.
Requires: pip install python-docx
Usage: python3 word_export.py
"""

from dataclasses import dataclass
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@dataclass
class Source:
    name: str
    url: str  # empty string -> flagged as unverifiable


@dataclass
class BodyRow:
    label: str
    values: list  # one entry per body_columns[1:]


@dataclass
class ResearchOutput:
    headline: str
    body_columns: list   # column headers including label column
    body_rows: list      # list of BodyRow
    bcg_opinion: str
    so_what: str
    sources: list        # list of Source


BCG_GREEN = RGBColor(0x00, 0x6A, 0x4E)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append a clickable hyperlink run to an existing paragraph."""
    r_id = paragraph.part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _shade_paragraph(paragraph, hex_fill: str) -> None:
    """Apply background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)


def export_to_word(output: ResearchOutput, filepath: str) -> str:
    """Write ResearchOutput to a formatted .docx and return the filepath."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 1. Headline
    h = doc.add_heading(output.headline, level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = BCG_GREEN

    doc.add_paragraph()

    # 2. Body table
    if output.body_rows and output.body_columns:
        table = doc.add_table(
            rows=1 + len(output.body_rows),
            cols=len(output.body_columns),
        )
        table.style = 'Light Shading'

        for i, col_name in enumerate(output.body_columns):
            cell = table.rows[0].cells[i]
            cell.text = col_name
            cell.paragraphs[0].runs[0].font.bold = True

        for r_idx, row in enumerate(output.body_rows):
            row_cells = table.rows[r_idx + 1].cells
            row_cells[0].text = row.label
            for c_idx, val in enumerate(row.values):
                if c_idx + 1 < len(row_cells):
                    row_cells[c_idx + 1].text = str(val)

    doc.add_paragraph()

    # 3. BCG 의견 (light green shaded box)
    opinion_p = doc.add_paragraph()
    opinion_p.paragraph_format.left_indent = Cm(0.5)
    opinion_p.paragraph_format.space_before = Pt(4)
    opinion_p.paragraph_format.space_after = Pt(4)
    _shade_paragraph(opinion_p, 'E8F5E9')

    label_run = opinion_p.add_run('[BCG 의견]  ')
    label_run.font.bold = True
    label_run.font.color.rgb = BCG_GREEN
    opinion_p.add_run(output.bcg_opinion)

    doc.add_paragraph()

    # 4. Source footer — name + clickable URL
    src_header = doc.add_paragraph()
    src_header.add_run('Source:').font.bold = True

    for src in output.sources:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f'{src.name}  —  ')
        if src.url:
            _add_hyperlink(p, src.url, src.url)
        else:
            flag = p.add_run('출처 불명확 — 추가 검증 필요')
            flag.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    bcg_p = doc.add_paragraph()
    bcg_p.paragraph_format.left_indent = Cm(0.5)
    bcg_p.add_run('BCG analysis')

    doc.add_paragraph()

    # 5. So-what footer (bold)
    sw = doc.add_paragraph()
    sw.add_run(f'특히, {output.so_what}').font.bold = True

    doc.save(filepath)
    return filepath


if __name__ == '__main__':
    result = export_to_word(
        ResearchOutput(
            headline='한국 건강기능식품 | TAM ~$2.1B, 온라인 채널 CAGR 18% 고성장 확인',
            body_columns=['구분', '정의', 'Q × P', '규모', '신뢰도'],
            body_rows=[
                BodyRow('TAM', ['전체 국내 건강기능식품', '5,200만 명 × $40/인', '$2.1B (약 2.9조 원)', 'H']),
                BodyRow('SAM', ['온라인 구매 가능 성인', '2,100만 명 × $40/인', '$840M (약 1.2조 원)', 'M']),
                BodyRow('SOM', ['3년 내 현실적 점유', 'SAM × 3%', '$25M (약 350억 원)', 'M']),
            ],
            bcg_opinion=(
                'SOM은 SAM의 3% 수준으로 KOL 채널 확보가 핵심 변수. '
                '온라인 집중 전략은 유효하나 오프라인 병행 없이는 재구매율 유지에 한계.'
            ),
            so_what='온라인 채널 진입 타당 → 따라서 KOL 파트너십 3개 이상 선확보 후 론칭 권고',
            sources=[
                Source('식품의약품안전처 건강기능식품 시장 현황 2024', 'https://www.mfds.go.kr'),
                Source('Nielsen Korea 온라인 채널 분석 2024', 'https://www.nielsen.com/kr'),
                Source('한국건강기능식품협회 연간 보고서', ''),  # no URL -- flagged
            ],
        ),
        '/tmp/bcg_research_output.docx',
    )
    print(f'Saved: {result}')
